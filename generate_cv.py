#!/usr/bin/env python3
"""Centralised CV document generator.

Generates a formatted .docx CV from a content dictionary.
Called by the cv-tailoring skill during document generation (Step 6).

Usage:
    from generate_cv import generate_cv
    generate_cv(content_dict, "/path/to/output.docx")
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

__all__ = ["generate_cv"]

# Formatting constants
DARK = RGBColor(0x2E, 0x34, 0x40)
LINK_BLUE = "0563C1"

# Required top-level keys and their expected types
_REQUIRED_FIELDS = {
    "name": str,
    "title": str,
    "contact": dict,
    "summary": list,
    "core_skills": str,
    "roles": list,
}

# Required keys within the contact sub-dict
_REQUIRED_CONTACT_FIELDS = {"email", "phone", "location"}


def _validate_content(content: dict) -> None:
    """Validate the content dict before generating the document.

    Raises:
        TypeError: If content is not a dict.
        ValueError: If required fields are missing or have the wrong type.
    """
    if not isinstance(content, dict):
        raise TypeError(f"content must be a dict, got {type(content).__name__}")

    for field, expected_type in _REQUIRED_FIELDS.items():
        if field not in content:
            raise ValueError(f"content is missing required field: '{field}'")
        if not isinstance(content[field], expected_type):
            raise ValueError(
                f"content['{field}'] must be {expected_type.__name__}, "
                f"got {type(content[field]).__name__}"
            )

    missing_contact = _REQUIRED_CONTACT_FIELDS - content["contact"].keys()
    if missing_contact:
        raise ValueError(
            f"content['contact'] is missing required fields: {sorted(missing_contact)}"
        )


def _safe_url(value: str) -> str:
    """Return a clean https:// URL from a value that may or may not include a protocol."""
    value = value.strip()
    if value.startswith("https://"):
        return value
    if value.startswith("http://"):
        return "https://" + value[7:]
    return f"https://{value}"


def _add_run(paragraph, text, bold=False, italic=False, size=None, color=DARK, font="Calibri"):
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.color.rgb = color
    if size:
        run.font.size = size
    return run


def _add_hyperlink(paragraph, text, url, size=Pt(9.5)):
    """Add a clickable hyperlink run to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    for tag, attr, val in [
        ("w:color", "w:val", LINK_BLUE),
        ("w:u", "w:val", "single"),
        ("w:sz", "w:val", str(int(size.pt * 2))),
    ]:
        el = OxmlElement(tag)
        el.set(qn(attr), val)
        rPr.append(el)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    rPr.append(rFonts)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _add_paragraph(doc, text="", bold=False, italic=False, size=None, color=DARK,
                   space_before=None, space_after=None):
    """Add a formatted paragraph to the document."""
    p = doc.add_paragraph()
    if text:
        _add_run(p, text, bold=bold, italic=italic, size=size, color=color)
    fmt = p.paragraph_format
    if space_before is not None:
        fmt.space_before = space_before
    if space_after is not None:
        fmt.space_after = space_after
    return p


def _add_section_heading(doc, text):
    """Add a section heading (Professional Summary, Experience, etc.)."""
    return _add_paragraph(doc, text, bold=True, size=Pt(11),
                          space_before=Pt(15), space_after=Pt(5))


def _add_bullet(doc, text):
    """Add a bullet-pointed paragraph using List Paragraph style."""
    p = doc.add_paragraph(text, style="List Paragraph")
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numPr.append(ilvl)
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), "1")
    numPr.append(numId)
    pPr.append(numPr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p


def _init_numbering(doc):
    """Ensure the numbering part exists by creating and removing a temp list paragraph."""
    temp = doc.add_paragraph("temp", style="List Bullet")
    temp._element.getparent().remove(temp._element)


def generate_cv(content: dict, output_path: str) -> str:
    """Generate a formatted CV .docx file.

    Args:
        content: Dict matching the CV content schema (see README or SKILL.md Step 6).
        output_path: Full path for the output .docx file. Parent directories
            are created automatically if they do not exist.

    Returns:
        The resolved output path as a string.

    Raises:
        TypeError: If content is not a dict.
        ValueError: If required fields are missing or have wrong types.
        OSError: If the output file cannot be written.
    """
    _validate_content(content)

    output = Path(output_path).resolve()
    output.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Page setup — A4
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.76)
    section.bottom_margin = Inches(0.69)
    section.left_margin = Inches(0.83)
    section.right_margin = Inches(0.83)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = DARK

    # Initialise numbering for bullet points
    _init_numbering(doc)

    # --- Header ---
    _add_paragraph(doc, content["name"], bold=True, size=Pt(18), space_after=Pt(2))
    _add_paragraph(doc, content["title"], size=Pt(12), space_after=Pt(1))

    # Contact line
    contact = content["contact"]
    p = _add_paragraph(doc, space_after=Pt(1))
    contact_text = f'{contact["email"]} | {contact["phone"]} | {contact["location"]}'
    _add_run(p, contact_text, size=Pt(9.5))

    # Links line: LinkedIn, GitHub, website, clearance — only rendered if provided
    linkedin = content.get("linkedin", "").strip()
    github = content.get("github", "").strip()
    website = content.get("website", "").strip()
    # additional_info covers any header note: clearance, work rights, nationality, etc.
    additional_info = (content.get("additional_info") or content.get("clearance") or "").strip()

    links = []
    if linkedin:
        links.append(("linkedin", linkedin))
    if github:
        links.append(("github", github))
    if website:
        links.append(("website", website))

    if links or additional_info:
        p = _add_paragraph(doc, space_after=Pt(5))
        for i, (label, value) in enumerate(links):
            if i > 0:
                _add_run(p, " | ", size=Pt(9.5))
            _add_hyperlink(p, value, _safe_url(value))
        if links and additional_info:
            _add_run(p, " | ", size=Pt(9.5))
        if additional_info:
            _add_run(p, additional_info, bold=True, size=Pt(9.5))

    # --- Professional Summary ---
    _add_section_heading(doc, "Professional Summary")
    for para_text in content["summary"]:
        _add_paragraph(doc, para_text, space_before=Pt(4), space_after=Pt(4))

    # --- Core Skills ---
    _add_section_heading(doc, "Core Skills")
    _add_paragraph(doc, content["core_skills"], size=Pt(9.5),
                   space_before=Pt(4), space_after=Pt(4))

    # --- Experience (Employment History) ---
    _add_section_heading(doc, "Experience")

    for role in content["roles"]:
        _add_paragraph(doc, role["title"], bold=True, size=Pt(10.5),
                       space_before=Pt(12), space_after=Pt(0))

        p = _add_paragraph(doc, space_before=Pt(1), space_after=Pt(1))
        _add_run(p, role["org"], italic=True)
        _add_run(p, f' | {role["dates"]}')

        if role.get("intro"):
            _add_paragraph(doc, role["intro"], space_before=Pt(4), space_after=Pt(4))

        for bullet in role.get("bullets", []):
            _add_bullet(doc, bullet)

    # --- Certifications ---
    if content.get("certifications"):
        _add_section_heading(doc, "Certifications")
        for cert in content["certifications"]:
            _add_paragraph(doc, f"\u2013  {cert}", space_after=Pt(2))

    # --- Courses ---
    if content.get("courses"):
        _add_section_heading(doc, "Courses")
        for course in content["courses"]:
            _add_paragraph(doc, f"\u2013  {course}", space_after=Pt(2))

    # --- Footer ---
    _add_paragraph(doc, "References available upon request.", italic=True,
                   space_before=Pt(12))

    doc.save(output)
    return str(output)
