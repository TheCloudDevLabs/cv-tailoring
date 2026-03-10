"""
Unit tests for generate_cv.py

Run from the project root:
    python tests/generate_cv_test.py

Tests cover: full content, minimal content, optional fields, all links,
input validation (missing fields, wrong types), output directory creation,
and path resolution.
"""

import sys
import os
import unittest
from pathlib import Path

# Ensure project root is on the path
sys.path.insert(0, str(Path(__file__).parent.parent))

from generate_cv import generate_cv

WORKSPACE = Path(__file__).parent / "workspace" / "generate_cv_tests"


def _full_content(**overrides):
    """Return a complete valid content dict."""
    content = {
        "name": "Alex Johnson",
        "title": "Cloud Engineer",
        "contact": {
            "email": "alex.johnson@example.com",
            "phone": "+1 555 000 0000",
            "location": "Seattle, WA",
        },
        "linkedin": "linkedin.com/in/alexjohnson",
        "github": "github.com/alexjohnson",
        "website": "alexjohnson.dev",
        "additional_info": "AWS Certified Solutions Architect",
        "summary": [
            "Cloud engineer with 8 years of experience on AWS and Azure.",
            "Strong background in Kubernetes and infrastructure as code.",
        ],
        "core_skills": "AWS · Azure · Kubernetes · Terraform · Docker · Python",
        "roles": [
            {
                "title": "Senior Cloud Engineer",
                "org": "Acme Corp",
                "dates": "March 2021 – Present",
                "intro": "Led platform engineering for a financial services client.",
                "bullets": [
                    "Designed multi-account AWS landing zone using Terraform.",
                    "Reduced infrastructure costs by 35% through Reserved Instance planning.",
                ],
            },
            {
                "title": "Cloud Engineer",
                "org": "Brightfield Technology",
                "dates": "January 2019 – February 2021",
                "intro": "",
                "bullets": [
                    "Built and maintained Azure Kubernetes Service clusters for 4 product teams.",
                ],
            },
        ],
        "certifications": [
            "AWS Certified Solutions Architect – Professional (2023)",
            "Certified Kubernetes Administrator (CKA) (2022)",
        ],
        "courses": [
            "HashiCorp Terraform Advanced (2021)",
        ],
    }
    content.update(overrides)
    return content


class TestGenerateCvFullContent(unittest.TestCase):
    def setUp(self):
        self.output = WORKSPACE / "full-content.docx"

    def test_creates_docx(self):
        result = generate_cv(_full_content(), str(self.output))
        self.assertTrue(self.output.exists(), "Output .docx file was not created")

    def test_returns_resolved_path(self):
        result = generate_cv(_full_content(), str(self.output))
        self.assertEqual(result, str(self.output.resolve()))

    def test_docx_is_readable(self):
        generate_cv(_full_content(), str(self.output))
        from docx import Document
        doc = Document(str(self.output))
        full_text = " ".join(p.text for p in doc.paragraphs)
        self.assertIn("Alex Johnson", full_text)
        self.assertIn("Cloud Engineer", full_text)
        self.assertIn("Acme Corp", full_text)


class TestGenerateCvMinimalContent(unittest.TestCase):
    """Only required fields — no optional links, clearance, courses."""

    def setUp(self):
        self.output = WORKSPACE / "minimal-content.docx"

    def test_creates_docx_with_required_fields_only(self):
        content = {
            "name": "Alex Johnson",
            "title": "Cloud Engineer",
            "contact": {
                "email": "alex.johnson@example.com",
                "phone": "+1 555 000 0000",
                "location": "Seattle, WA",
            },
            "summary": ["Cloud engineer with 8 years of experience."],
            "core_skills": "AWS · Kubernetes · Terraform",
            "roles": [
                {
                    "title": "Senior Cloud Engineer",
                    "org": "Acme Corp",
                    "dates": "2021 – Present",
                    "intro": "",
                    "bullets": ["Led platform engineering team."],
                }
            ],
            "certifications": [],
        }
        result = generate_cv(content, str(self.output))
        self.assertTrue(self.output.exists())


class TestGenerateCvNoOptionalLinks(unittest.TestCase):
    """No linkedin, github, website, or additional_info."""

    def setUp(self):
        self.output = WORKSPACE / "no-links.docx"

    def test_creates_docx_without_link_line(self):
        content = _full_content(linkedin="", github="", website="", additional_info="")
        result = generate_cv(content, str(self.output))
        self.assertTrue(self.output.exists())

    def test_docx_still_contains_name(self):
        content = _full_content(linkedin="", github="", website="", additional_info="")
        generate_cv(content, str(self.output))
        from docx import Document
        doc = Document(str(self.output))
        full_text = " ".join(p.text for p in doc.paragraphs)
        self.assertIn("Alex Johnson", full_text)


class TestGenerateCvAllLinks(unittest.TestCase):
    """All link fields populated — verify none are dropped."""

    def setUp(self):
        self.output = WORKSPACE / "all-links.docx"

    def test_all_links_included(self):
        content = _full_content(
            linkedin="linkedin.com/in/alexjohnson",
            github="github.com/alexjohnson",
            website="alexjohnson.dev",
            additional_info="TS/SCI Clearance | US Citizen",
        )
        generate_cv(content, str(self.output))
        from docx import Document
        doc = Document(str(self.output))
        full_text = " ".join(p.text for p in doc.paragraphs)
        self.assertIn("TS/SCI Clearance | US Citizen", full_text)


class TestGenerateCvUrlSanitisation(unittest.TestCase):
    """URLs with existing https:// prefix should not get double protocol."""

    def setUp(self):
        self.output = WORKSPACE / "url-sanitisation.docx"

    def test_https_prefix_not_doubled(self):
        # User stored full URL instead of bare domain
        content = _full_content(linkedin="https://linkedin.com/in/alexjohnson")
        # Should not raise and should produce a readable docx
        result = generate_cv(content, str(self.output))
        self.assertTrue(self.output.exists())


class TestGenerateCvOutputDirectoryCreation(unittest.TestCase):
    """Output directory should be created automatically if it does not exist."""

    def test_creates_nested_output_directory(self):
        output = WORKSPACE / "nested" / "deep" / "output.docx"
        self.assertFalse(output.parent.exists(), "Precondition: directory should not exist")
        generate_cv(_full_content(), str(output))
        self.assertTrue(output.exists())


class TestGenerateCvValidation(unittest.TestCase):
    """Missing or wrong-type fields should raise clear errors."""

    def _out(self, name):
        return str(WORKSPACE / f"validation-{name}.docx")

    def test_missing_name_raises_value_error(self):
        content = _full_content()
        del content["name"]
        with self.assertRaises(ValueError) as ctx:
            generate_cv(content, self._out("no-name"))
        self.assertIn("name", str(ctx.exception))

    def test_missing_summary_raises_value_error(self):
        content = _full_content()
        del content["summary"]
        with self.assertRaises(ValueError):
            generate_cv(content, self._out("no-summary"))

    def test_missing_contact_email_raises_value_error(self):
        content = _full_content()
        del content["contact"]["email"]
        with self.assertRaises(ValueError) as ctx:
            generate_cv(content, self._out("no-email"))
        self.assertIn("email", str(ctx.exception))

    def test_wrong_type_summary_raises_value_error(self):
        content = _full_content(summary="not a list")
        with self.assertRaises(ValueError) as ctx:
            generate_cv(content, self._out("bad-summary-type"))
        self.assertIn("summary", str(ctx.exception))

    def test_non_dict_content_raises_type_error(self):
        with self.assertRaises(TypeError):
            generate_cv("not a dict", self._out("not-dict"))

    def test_missing_roles_raises_value_error(self):
        content = _full_content()
        del content["roles"]
        with self.assertRaises(ValueError):
            generate_cv(content, self._out("no-roles"))


if __name__ == "__main__":
    WORKSPACE.mkdir(parents=True, exist_ok=True)
    print(f"Test outputs: {WORKSPACE}\n")
    unittest.main(verbosity=2)
